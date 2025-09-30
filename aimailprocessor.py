#!/usr/bin/env python3
# coding: utf-8
"""
autor: technoconsulx
"""

import asyncio
import base64
import email
import json
import logging
import os
import re
import ssl
import sys
from datetime import datetime
from email import policy
from email.parser import BytesParser
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
import imaplib
import requests
from aiosmtplib import SMTP
from bs4 import BeautifulSoup

# Try imports for extractors
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx as python_docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    python_docx = None

try:
    import textract
except Exception:
    textract = None

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
except Exception:
    openpyxl = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    # Попробуем использовать стандартные шрифты вместо Times New Roman
    from reportlab.pdfgen.canvas import Canvas
    from reportlab.lib.units import inch
except Exception:
    SimpleDocTemplate = None

# ==========================
# CONFIGURATION
# ==========================
IMAP_HOST = "*****"
IMAP_USER = "*****"
IMAP_PASS = "*****"
SMTP_HOST = "*****"
SMTP_PORT = 25
SMTP_USER = "*****"
SMTP_PASS = "*****"
CHECK_INTERVAL = 30  # seconds
TELEGRAMS_DIR = "telegrams"
ERROR_DIR = "ErrorTelegrams"
PROCESSED_DIR = "ProcessedTelegrams"
OUTPUT_DIR = "GeneratedDocuments"
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "gemma3:12b"
OLLAMA_TIMEOUT = 9999  # seconds
DUCKDUCKGO_MAX = 3  # number of search results to include as context
MAX_ATTACHMENT_SIZE = 10 * 1024 * 1024  # 10 MB

# Create required dirs
os.makedirs(TELEGRAMS_DIR, exist_ok=True)
os.makedirs(ERROR_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ==========================
# LOGGING
# ==========================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# ==========================
# Utility: safe filename + save bytes
# ==========================
def safe_filename(name: str) -> str:
    return re.sub(r"[^\w\-. ]", "_", name)

def save_bytes_to_file(dirname: str, filename: str, data: bytes) -> str:
    if len(data) > MAX_ATTACHMENT_SIZE:
        logging.warning(f"Skipping {filename}: exceeds 10 MB")
        return None
    
    fname = safe_filename(filename)
    path = os.path.join(dirname, fname)
    base, ext = os.path.splitext(path)
    counter = 1
    while os.path.exists(path):
        path = f"{base}_{counter}{ext}"
        counter += 1
    
    with open(path, "wb") as f:
        f.write(data)
    return path

# ==========================
# Document Creation Functions (с исправленной кодировкой PDF)
# ==========================

def create_official_word_document(ai_response: str, filename: str, is_table_request: bool = False) -> str:
    """Создает официальный Word документ в деловом стиле"""
    if not python_docx:
        logging.error("python-docx not available")
        return None
    
    try:
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        
        # Заголовок документа
        title = doc.add_heading('', level=1)
        title_run = title.add_run('Ответ AI Ассистента')
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата создания
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        date_run.italic = True
        
        doc.add_paragraph()  # Пустая строка
        
        # Очищаем ответ от лишних символов
        clean_response = clean_ai_response(ai_response)
        
        # Основное содержимое
        if is_table_request and has_table_data(clean_response):
            # Для табличных данных создаем таблицу
            table_data = parse_table_data(clean_response)
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = str(cell)
                        table.cell(i, j).paragraphs[0].runs[0].font.name = 'Times New Roman'
                        table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(12)
        else:
            # Обычный текст с форматированием
            paragraphs = clean_response.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style = doc.styles['Normal']
        
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logging.info(f"Created official Word document: {filepath}")
        return filepath
    except Exception as e:
        logging.error(f"Failed to create official Word document: {e}")
        return None

def create_simple_pdf_fallback(ai_response: str, filename: str, is_table_request: bool = False) -> str:
    """Простой fallback метод для создания PDF когда основной не работает"""
    try:
        filepath = os.path.join(OUTPUT_DIR, filename)
        c = Canvas(filepath, pagesize=A4)
        width, height = A4
        
        # Используем стандартные шрифты, которые точно есть
        # Пробуем Helvetica, который имеет базовую поддержку кириллицы в некоторых системах
        c.setFont("Helvetica", 12)
        y_position = height - 50
        line_height = 14
        
        # Заголовок
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y_position, "Ответ AI Ассистента")
        y_position -= 30
        
        # Дата
        c.setFont("Helvetica-Oblique", 10)
        c.drawString(50, y_position, f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        y_position -= 30
        
        c.setFont("Helvetica", 12)
        
        clean_response = clean_ai_response(ai_response)
        
        if is_table_request and has_table_data(clean_response):
            # Простая таблица
            table_data = parse_table_data(clean_response)
            if table_data:
                col_width = (width - 100) / len(table_data[0])
                
                for i, row in enumerate(table_data):
                    if y_position < 100:  # Новая страница
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y_position = height - 50
                    
                    x_position = 50
                    for j, cell in enumerate(row):
                        c.drawString(x_position, y_position, str(cell)[:30])  # Обрезаем длинный текст
                        x_position += col_width
                    
                    y_position -= line_height
                    if i == 0:  # Линия после заголовка
                        y_position -= 5
        else:
            # Обычный текст
            paragraphs = clean_response.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    lines = split_text_to_lines(paragraph, 80)
                    for line in lines:
                        if y_position < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y_position = height - 50
                        
                        c.drawString(50, y_position, line)
                        y_position -= line_height
                    
                    y_position -= 5
        
        c.save()
        logging.info(f"Created fallback PDF document: {filepath}")
        return filepath
        
    except Exception as e:
        logging.error(f"Fallback PDF creation failed: {e}")
        return None

def create_official_pdf_document(ai_response: str, filename: str, is_table_request: bool = False) -> str:
    """Создает официальный PDF документ с исправленной кодировкой и поддержкой русского языка"""
    if not SimpleDocTemplate:
        logging.error("reportlab not available")
        return None
    
    try:
        filepath = os.path.join(OUTPUT_DIR, filename)
        
        # Регистрируем шрифты с поддержкой кириллицы
        try:
            # Попробуем использовать DejaVu Sans - обычно установлен в системе
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
            main_font = 'DejaVuSans'
            bold_font = 'DejaVuSans-Bold'
        except:
            try:
                # Попробуем Liberation Sans
                pdfmetrics.registerFont(TTFont('LiberationSans', '/usr/share/fonts/liberation/LiberationSans-Regular.ttf'))
                pdfmetrics.registerFont(TTFont('LiberationSans-Bold', '/usr/share/fonts/liberation/LiberationSans-Bold.ttf'))
                main_font = 'LiberationSans'
                bold_font = 'LiberationSans-Bold'
            except:
                try:
                    # FreeSans как запасной вариант
                    pdfmetrics.registerFont(TTFont('FreeSans', '/usr/share/fonts/truetype/freefont/FreeSans.ttf'))
                    pdfmetrics.registerFont(TTFont('FreeSans-Bold', '/usr/share/fonts/truetype/freefont/FreeSansBold.ttf'))
                    main_font = 'FreeSans'
                    bold_font = 'FreeSans-Bold'
                except:
                    # Если шрифты не найдены, используем стандартные (могут быть проблемы с кириллицей)
                    logging.warning("Cyrillic fonts not found, using default fonts (may display squares)")
                    main_font = 'Helvetica'
                    bold_font = 'Helvetica-Bold'
        
        # Используем SimpleDocTemplate для лучшего форматирования
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Создаем стили с правильными шрифтами
        styles = getSampleStyleSheet()
        
        # Основной стиль с русским шрифтом
        normal_style = ParagraphStyle(
            'NormalRU',
            parent=styles['Normal'],
            fontName=main_font,
            fontSize=12,
            leading=14,
            spaceAfter=12
        )
        
        # Стиль для заголовка
        title_style = ParagraphStyle(
            'TitleRU', 
            parent=styles['Heading1'],
            fontName=bold_font,
            fontSize=16,
            alignment=1,  # CENTER
            spaceAfter=30
        )
        
        # Стиль для даты
        date_style = ParagraphStyle(
            'DateRU',
            parent=styles['Normal'],
            fontName=main_font,
            fontSize=10,
            alignment=2,  # RIGHT
            spaceAfter=30,
            textColor=colors.gray
        )
        
        # Собираем содержимое документа
        story = []
        
        # Заголовок
        title = Paragraph("Ответ AI Ассистента", title_style)
        story.append(title)
        
        # Дата создания
        date_text = f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        date_para = Paragraph(date_text, date_style)
        story.append(date_para)
        
        # Пустой отступ
        story.append(Spacer(1, 20))
        
        # Очищаем ответ
        clean_response = clean_ai_response(ai_response)
        
        # Обрабатываем содержимое
        if is_table_request and has_table_data(clean_response):
            # Создаем таблицу
            table_data = parse_table_data(clean_response)
            if table_data:
                # Создаем таблицу с данными
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), main_font, 10),
                    ('FONT', (0, 0), (-1, 0), bold_font, 10),  # Заголовки жирным
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
        else:
            # Обычный текст
            paragraphs = clean_response.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Заменяем множественные пробелы на одинарные
                    clean_para = re.sub(r'\s+', ' ', paragraph.strip())
                    para = Paragraph(clean_para, normal_style)
                    story.append(para)
        
        # Собираем документ
        doc.build(story)
        logging.info(f"Created official PDF document with Cyrillic support: {filepath}")
        return filepath
        
    except Exception as e:
        logging.error(f"Failed to create official PDF document: {e}")
        
        # Fallback: простой метод с Canvas если основной не сработал
        try:
            return create_simple_pdf_fallback(ai_response, filename, is_table_request)
        except Exception as e2:
            logging.error(f"Fallback PDF creation also failed: {e2}")
            return None

def create_official_excel_document(ai_response: str, filename: str, is_table_request: bool = False) -> str:
    """Создает официальный Excel документ"""
    if not openpyxl:
        logging.error("openpyxl not available")
        return None
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Документ"
        
        # Настройка стилей
        header_font = Font(name='Times New Roman', size=12, bold=True)
        normal_font = Font(name='Times New Roman', size=11)
        thin_border = Border(left=Side(style='thin'), 
                           right=Side(style='thin'), 
                           top=Side(style='thin'), 
                           bottom=Side(style='thin'))
        
        current_row = 1
        
        # Очищаем ответ от лишних символов
        clean_response = clean_ai_response(ai_response)
        
        if is_table_request and has_table_data(clean_response):
            # Для табличных данных создаем таблицу
            table_data = parse_table_data(clean_response)
            if table_data:
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        ws.cell(row=current_row + i, column=j + 1, value=str(cell))
                        ws.cell(row=current_row + i, column=j + 1).font = normal_font
                        ws.cell(row=current_row + i, column=j + 1).border = thin_border
                        if i == 0:  # Заголовки
                            ws.cell(row=current_row + i, column=j + 1).font = header_font
                
                # Автоподбор ширины колонок
                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    ws.column_dimensions[column].width = adjusted_width
        else:
            # Обычный текст - каждая строка в отдельной ячейке
            paragraphs = clean_response.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    ws.cell(row=current_row, column=1, value=paragraph.strip())
                    ws.cell(row=current_row, column=1).font = normal_font
                    current_row += 1
            
            # Автоподбор ширины для текста
            ws.column_dimensions['A'].width = 50
        
        filepath = os.path.join(OUTPUT_DIR, filename)
        wb.save(filepath)
        logging.info(f"Created official Excel document: {filepath}")
        return filepath
    except Exception as e:
        logging.error(f"Failed to create official Excel document: {e}")
        return None

def clean_ai_response(text: str) -> str:
    """Очищает ответ AI от лишних символов и форматирования"""
    if not text:
        return ""
    
    # Убираем звездочки и другие спецсимволы
    text = re.sub(r'\*+', '', text)  # Убирает ***
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Убирает **жирный**
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Убирает *курсив*
    
    # Убираем лишние пробелы и переносы
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Убирает множественные пустые строки
    text = re.sub(r'[ \t]+', ' ', text)  # Убирает множественные пробелы/табы
    
    # Убираем маркеры документов если остались
    text = re.sub(r'\[CREATE_DOCUMENTS:[^\]]+\]', '', text)
    
    # Убираем лишние символы в начале и конце
    text = text.strip()
    
    return text

def split_text_to_lines(text: str, max_line_length: int) -> list:
    """Разбивает текст на строки указанной максимальной длины"""
    words = text.split()
    lines = []
    current_line = []
    
    for word in words:
        if len(' '.join(current_line + [word])) <= max_line_length:
            current_line.append(word)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
    
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines

def has_table_data(text: str) -> bool:
    """Определяет, содержит ли текст табличные данные"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    if len(lines) < 2:
        return False
    
    # Проверяем признаки таблицы
    table_indicators = [
        any(sep in line for sep in ['|', '\t']) for line in lines
    ]
    
    # Если больше половины строк содержат табличные признаки
    return sum(table_indicators) > len(lines) * 0.3

def parse_table_data(text: str) -> list:
    """Парсит табличные данные из текста"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    table_data = []
    
    for line in lines:
        if '|' in line:
            # Разделяем по вертикальным чертам
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                table_data.append(cells)
        elif '\t' in line:
            # Разделяем по табуляции
            cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
            if cells:
                table_data.append(cells)
        elif re.match(r'.+\s{2,}.+', line):
            # Разделяем по двум и более пробелам
            cells = re.split(r'\s{2,}', line.strip())
            if cells:
                table_data.append(cells)
    
    return table_data if len(table_data) > 1 else None

# ==========================
# Enhanced Ollama call with improved prompt
# ==========================
def call_ollama_with_smart_document_detection(prompt: str, image_paths: list = None, extra_ctx: str = None):
    images_b64 = []
    if image_paths:
        for p in image_paths:
            try:
                with open(p, "rb") as f:
                    images_b64.append(base64.b64encode(f.read()).decode("utf-8"))
            except Exception as e:
                logging.warning(f"Failed to encode image {p}: {e}")
    
    # Улучшенный системный промпт для предотвращения спама символами
    system_prompt = """Ты - дружелюбный AI-ассистент, который отвечает на входящие письма. 

ВНИМАНИЕ: НЕ используй символы * для форматирования! Отвечай чистым текстом.

АНАЛИЗИРУЙ запрос пользователя и определи:
1. Если просят ТАБЛИЦУ, СПИСОК, EXCEL - используй табличный формат с | разделителями
2. Если просят ДОКУМЕНТ, ОТЧЕТ, WORD, PDF - отвечай обычным текстом без таблиц
3. Если не ясно - используй обычный текст

ТАБЛИЦУ используй ТОЛЬКО когда явно просят:
- "создай таблицу", "сделай список", "в виде таблицы", "excel"
- для перечней, сравнений, структурированных данных

ОБЫЧНЫЙ ТЕКСТ для всех остальных случаев:
- документы, отчеты, письма, описания

ФОРМАТИРОВАНИЕ:
- НЕ используй * ** для выделения текста
- НЕ используй markdown разметку
- Используй обычные заголовки и абзацы

ЕСЛИ нужны документы - в конце ответа добавь маркер и укажи тип:
[CREATE_DOCUMENTS:docx,pdf,xlsx&type=table] - для таблиц
[CREATE_DOCUMENTS:docx,pdf,xlsx&type=document] - для документов

Отвечай естественно и по делу, без лишних символов форматирования."""

    full_prompt = (extra_ctx + "\n\n" if extra_ctx else "") + (prompt or "")
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": full_prompt,
        "stream": False,
        "temperature": 0.7,
        "top_p": 0.9,
        "num_ctx": 16384,
        "system": system_prompt,
        "num_predict": 6000
    }
    
    if images_b64:
        payload["images"] = images_b64
    
    try:
        logging.info(f"Calling Ollama for smart document detection")
        resp = requests.post(OLLAMA_URL, json=payload, timeout=OLLAMA_TIMEOUT)
        resp.raise_for_status()
        return resp.json()
    except requests.exceptions.RequestException as e:
        logging.error(f"Ollama request failed: {e}")
        return None

def extract_document_requirements(response_text: str):
    """Извлекает требования к документам из ответа ИИ"""
    doc_pattern = r'\[CREATE_DOCUMENTS:([^\]]+)\]'
    match = re.search(doc_pattern, response_text)
    
    if match:
        doc_info = match.group(1)
        formats = doc_info.split('&')[0].split(',')
        
        # Определяем тип документа
        doc_type = "document"  # по умолчанию обычный документ
        if 'type=table' in doc_info:
            doc_type = "table"
        elif 'type=document' in doc_info:
            doc_type = "document"
        
        # Убираем маркер из основного текста
        clean_text = re.sub(doc_pattern, '', response_text).strip()
        return clean_text, formats, doc_type
    
    return response_text, [], "document"

def is_table_request(text: str) -> bool:
    """Определяет, является ли запрос запросом на таблицу"""
    table_keywords = [
        'таблиц', 'список', 'excel', 'табличн', 'перечень', 
        'сетка', 'матриц', 'сравнен', 'рейтинг'
    ]
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in table_keywords)

# ==========================
# Extract attachments and text (остается без изменений)
# ==========================
def extract_attachments_and_text(msg) -> (list, str):
    attachments = []
    text_parts = []
    
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            disp = str(part.get("Content-Disposition") or "")
            
            if ctype == "text/plain" and "attachment" not in disp:
                try:
                    raw = part.get_payload(decode=True)
                    if raw:
                        charset = part.get_content_charset() or "utf-8"
                        text = raw.decode(charset, errors="ignore")
                        text_parts.append(text)
                except Exception as e:
                    logging.warning(f"Failed to decode text/plain part: {e}")
            
            filename = part.get_filename()
            if filename:
                try:
                    payload = part.get_payload(decode=True) or b""
                    path = save_bytes_to_file(TELEGRAMS_DIR, filename, payload)
                    if path:
                        attachments.append(path)
                        logging.info(f"Saved attachment: {path}")
                except Exception as e:
                    logging.error(f"Failed to save attachment {filename}: {e}")
    else:
        try:
            raw = msg.get_payload(decode=True)
            if raw:
                charset = msg.get_content_charset() or "utf-8"
                try:
                    text_parts.append(raw.decode(charset, errors="ignore"))
                except Exception:
                    text_parts.append(raw.decode("utf-8", errors="ignore"))
        except Exception as e:
            logging.warning(f"Failed to extract singlepart body: {e}")

    # Try to extract text from supported attachments
    extracted_texts = []
    for path in attachments:
        if not path:
            continue
        lower = path.lower()
        try:
            if lower.endswith(".pdf") and PyPDF2:
                try:
                    with open(path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        pages = []
                        for page in reader.pages:
                            try:
                                page_text = page.extract_text() or ""
                                pages.append(page_text)
                            except Exception:
                                pass
                        extracted_texts.append("\n".join(pages))
                    logging.info(f"Extracted text from PDF: {path}")
                except Exception as e:
                    logging.warning(f"PDF extraction failed for {path}: {e}")
            
            elif lower.endswith(".docx") and python_docx:
                try:
                    doc = python_docx.Document(path)
                    texts = [p.text for p in doc.paragraphs if p.text]
                    extracted_texts.append("\n".join(texts))
                    logging.info(f"Extracted text from DOCX: {path}")
                except Exception as e:
                    logging.warning(f"DOCX extraction failed for {path}: {e}")
            
            elif lower.endswith(".doc"):
                if textract:
                    try:
                        raw = textract.process(path)
                        extracted_texts.append(raw.decode("utf-8", errors="ignore"))
                        logging.info(f"Extracted text from DOC via textract: {path}")
                    except Exception as e:
                        logging.warning(f"textract failed for {path}: {e}")
                else:
                    logging.warning(f"Skipping .doc extraction (textract not available): {path}")
            
            elif lower.endswith(".txt"):
                try:
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        extracted_texts.append(f.read())
                    logging.info(f"Read txt file: {path}")
                except Exception as e:
                    logging.warning(f"TXT read failed: {e}")
            
            elif lower.endswith(".eml"):
                try:
                    with open(path, "rb") as f:
                        em = BytesParser(policy=policy.default).parsebytes(f.read())
                    parts = []
                    if em.is_multipart():
                        for p in em.walk():
                            if p.get_content_type() == "text/plain":
                                payload = p.get_payload(decode=True)
                                if payload:
                                    parts.append(payload.decode(p.get_content_charset() or "utf-8", errors="ignore"))
                    else:
                        payload = em.get_payload(decode=True)
                        if payload:
                            parts.append(payload.decode(em.get_content_charset() or "utf-8", errors="ignore"))
                    extracted_texts.append("\n".join(parts))
                    logging.info(f"Extracted text from EML: {path}")
                except Exception as e:
                    logging.warning(f"EML extraction failed: {e}")
            
            elif lower.endswith(".xlsx") and openpyxl:
                try:
                    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                    sheets_text = []
                    for sheet in wb.worksheets:
                        rows = []
                        for row in sheet.iter_rows(values_only=True):
                            cells = [str(c) for c in row if c is not None]
                            if cells:
                                rows.append("\t".join(cells))
                        if rows:
                            sheets_text.append("\n".join(rows))
                    extracted_texts.append("\n\n".join(sheets_text))
                    logging.info(f"Extracted text from XLSX: {path}")
                except Exception as e:
                    logging.warning(f"XLSX extraction failed for {path}: {e}")
            else:
                logging.debug(f"No extractor for {path}")
        except Exception as e:
            logging.warning(f"Error extracting from {path}: {e}")

    combined_text = "\n\n".join([t for t in (text_parts + extracted_texts) if t])
    return attachments, combined_text

# ==========================
# DuckDuckGo quick search (scrape) - без изменений
# ==========================
def duckduckgo_search(query: str, max_results: int = DUCKDUCKGO_MAX) -> list:
    try:
        params = {"q": query}
        headers = {"User-Agent": "Mozilla/5.0 (compatible; bot/1.0)"}
        r = requests.get("https://duckduckgo.com/html/", params=params, headers=headers, timeout=15)
        r.raise_for_status()
        
        soup = BeautifulSoup(r.text, "lxml")
        results = []
        
        for a in soup.select("a.result__a")[:max_results]:
            title = a.get_text(strip=True)
            href = a.get("href")
            parent = a.find_parent("div", class_="result")
            snippet = ""
            if parent:
                s = parent.select_one(".result__snippet")
                if s:
                    snippet = s.get_text(strip=True)
            results.append((title, href, snippet))
        
        if not results:
            for rdiv in soup.select(".result")[:max_results]:
                a = rdiv.find("a", href=True)
                if a:
                    title = a.get_text(strip=True)
                    href = a["href"]
                    s = rdiv.find(class_="result__snippet")
                    snippet = s.get_text(strip=True) if s else ""
                    results.append((title, href, snippet))
        
        return results[:max_results]
    except Exception as e:
        logging.warning(f"DuckDuckGo search failed: {e}")
        return []

# ==========================
# SMTP send with attachments - без изменений
# ==========================
async def send_reply(to_addr: str, subject: str, body: str, attachments: list = None):
    try:
        # Создаем multipart сообщение для вложений
        if attachments:
            msg = MIMEMultipart()
            msg.attach(MIMEText(body, "plain", "utf-8"))
            
            for filepath in attachments:
                try:
                    with open(filepath, "rb") as f:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                        email.encoders.encode_base64(part)
                        part.add_header(
                            'Content-Disposition',
                            f'attachment; filename="{os.path.basename(filepath)}"'
                        )
                        msg.attach(part)
                except Exception as e:
                    logging.error(f"Failed to attach {filepath}: {e}")
        else:
            msg = MIMEText(body, "plain", "utf-8")
        
        msg["Subject"] = f"Re: {subject or 'Без темы'}"
        msg["From"] = SMTP_USER
        msg["To"] = to_addr
        msg["Reply-To"] = SMTP_USER

        smtp = SMTP(
            hostname=SMTP_HOST,
            port=SMTP_PORT,
            use_tls=False,
            start_tls=False,
            tls_context=None
        )

        await smtp.connect()
        await smtp.login(SMTP_USER, SMTP_PASS)
        await smtp.send_message(msg)
        await smtp.quit()
        
        logging.info(f"Reply sent to {to_addr} with {len(attachments or [])} attachments")
        
    except Exception as e:
        logging.error(f"send_reply failed: {e}")
        logging.error(f"SMTP details: host={SMTP_HOST}, port={SMTP_PORT}, user={SMTP_USER}, to={to_addr}")
        raise

# ==========================
# IMAP helpers - без изменений
# ==========================
def fetch_unseen_uids(mail) -> list:
    typ, data = mail.search(None, "UNSEEN")
    if typ != "OK":
        logging.warning("IMAP search failed")
        return []
    return data[0].split()

def fetch_full_message(mail, uid_bytes):
    typ, data = mail.fetch(uid_bytes, "(RFC822)")
    if typ != "OK":
        raise RuntimeError("IMAP fetch failed")
    return data[0][1]

# ==========================
# Smart message processor - без изменений
# ==========================
async def handle_message(mail, uid_bytes):
    raw = fetch_full_message(mail, uid_bytes)
    msg = BytesParser(policy=policy.default).parsebytes(raw)
    
    subject = msg.get("subject", "(no subject)")
    from_header = msg.get("from", "")
    
    from_addr = ""
    if from_header:
        email_match = re.search(r'<([^>]+)>', from_header)
        if email_match:
            from_addr = email_match.group(1)
        else:
            from_addr = from_header.strip()
    
    if not from_addr:
        logging.error(f"Cannot extract email from: {from_header}")
        return
    
    uid_str = uid_bytes.decode() if isinstance(uid_bytes, bytes) else str(uid_bytes)
    
    logging.info(f"Handling msg uid={uid_str} subject={subject} from={from_addr}")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    raw_path = os.path.join(TELEGRAMS_DIR, f"raw_{timestamp}_{uid_str}.eml")
    
    try:
        with open(raw_path, "wb") as f:
            f.write(raw)
    except Exception as e:
        logging.warning(f"Failed to save raw email: {e}")
    
    attachments, combined_text = extract_attachments_and_text(msg)
    logging.info(f"Found {len(attachments)} attachments; extracted_text_len={len(combined_text)}")
    
    extra_ctx = ""
    if not combined_text or len(combined_text.strip()) < 80:
        q_parts = [subject or "", from_addr or ""]
        if attachments:
            q_parts.append(os.path.basename(attachments[0]))
        query = " ".join([p for p in q_parts if p]).strip()
        
        if query:
            dd = duckduckgo_search(query, max_results=DUCKDUCKGO_MAX)
            if dd:
                lines = ["DuckDuckGo quick search results:"]
                for title, url, snip in dd:
                    lines.append(f"- {title} | {url}\n  {snip}")
                extra_ctx = "\n".join(lines)
                logging.info(f"Added DuckDuckGo context ({len(dd)} results) for uid={uid_str}")
    
    filenames = [os.path.basename(p) for p in attachments if p]
    prompt_parts = []
    
    if combined_text:
        prompt_parts.append("Письмо / извлечённый текст:\n" + combined_text)
    if filenames:
        prompt_parts.append("Вложенные файлы: " + ", ".join(filenames))
    
    prompt = "\n\n".join(prompt_parts) if prompt_parts else "(No text provided)"
    image_paths = [p for p in attachments if p and p.lower().endswith((".jpg", ".jpeg", ".png", ".webp"))]
    
    # Определяем тип запроса
    is_table_req = is_table_request(combined_text)
    
    # Используем умный вызов ИИ
    response = call_ollama_with_smart_document_detection(prompt, image_paths=image_paths, extra_ctx=extra_ctx)
    
    response_path = os.path.join(TELEGRAMS_DIR, f"response_{timestamp}_{uid_str}.txt")
    
    body = ""
    document_formats = []
    document_type = "document"
    
    if isinstance(response, dict):
        response_text = ""
        for k in ("response", "generated_text", "text", "output", "result"):
            if k in response and isinstance(response[k], str) and response[k].strip():
                response_text = response[k]
                break
        
        if not response_text:
            response_text = json.dumps(response, ensure_ascii=False, indent=2)
        
        # Извлекаем требования к документам
        clean_body, document_formats, doc_type_from_ai = extract_document_requirements(response_text)
        body = clean_body
        document_type = doc_type_from_ai
        
        # Если ИИ не определил тип, используем наш анализ
        if document_type == "document" and is_table_req:
            document_type = "table"
    
    else:
        body = str(response)
        document_type = "table" if is_table_req else "document"
    
    if not body:
        body = "(No response from model)"
    
    # Очищаем тело письма от лишних символов
    body = clean_ai_response(body)
    
    # Сохраняем полный ответ
    try:
        full_response_data = {
            "original_response": response,
            "cleaned_body": body,
            "document_formats": document_formats,
            "document_type": document_type,
            "is_table_request": is_table_req,
            "timestamp": timestamp
        }
        with open(response_path, "w", encoding="utf-8") as f:
            f.write(json.dumps(full_response_data, ensure_ascii=False, indent=2))
        logging.info(f"Saved enhanced response to {response_path}")
    except Exception as e:
        logging.error(f"Failed to save response: {e}")
    
    # Генерация документов на основе типа
    generated_attachments = []
    if document_formats:
        logging.info(f"AI requested {document_type} generation in formats: {document_formats}")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"document_{timestamp}_{uid_str}"
        
        is_table_doc = (document_type == "table")
        
        if 'docx' in document_formats or 'word' in document_formats:
            docx_file = create_official_word_document(body, f"{base_name}.docx", is_table_doc)
            if docx_file:
                generated_attachments.append(docx_file)
        
        if 'pdf' in document_formats:
            pdf_file = create_official_pdf_document(body, f"{base_name}.pdf", is_table_doc)
            if pdf_file:
                generated_attachments.append(pdf_file)
        
        if 'xlsx' in document_formats or 'excel' in document_formats:
            excel_file = create_official_excel_document(body, f"{base_name}.xlsx", is_table_doc)
            if excel_file:
                generated_attachments.append(excel_file)
    
    try:
        await send_reply(from_addr, subject, body, generated_attachments)
        try:
            processed_path = os.path.join(PROCESSED_DIR, os.path.basename(raw_path))
            os.replace(raw_path, processed_path)
        except Exception:
            pass
    except Exception as e:
        logging.error(f"Failed to send reply for uid={uid_str}: {e}")
        err_path = os.path.join(ERROR_DIR, f"error_{timestamp}_{uid_str}.eml")
        try:
            with open(err_path, "wb") as f:
                f.write(raw)
            logging.error(f"Moved problematic email to {err_path}")
        except Exception as e2:
            logging.error(f"Also failed to save error eml: {e2}")

# ==========================
# MAIN LOOP - без изменений
# ==========================
async def main_loop():
    logging.info("Starting Mail AI Processor with Fixed PDF Encoding")
    
    while True:
        try:
            logging.info("Connecting to IMAP...")
            mail = imaplib.IMAP4_SSL(IMAP_HOST)
            mail.login(IMAP_USER, IMAP_PASS)
            mail.select("inbox")
            
            uids = fetch_unseen_uids(mail)
            logging.info(f"Found {len(uids)} new messages")
            
            for uid in uids:
                try:
                    await handle_message(mail, uid)
                except Exception as e:
                    logging.error(f"Error handling uid {uid}: {e}")
            
            mail.logout()
            logging.info("IMAP logout")
            
        except Exception as e:
            logging.error(f"Main loop error: {e}")
        
        logging.info(f"Sleeping {CHECK_INTERVAL} seconds until next check")
        await asyncio.sleep(CHECK_INTERVAL)

if __name__ == "__main__":
    try:
        asyncio.run(main_loop())
    except KeyboardInterrupt:
        logging.info("Interrupted by user, exiting")
    except Exception as e:
        logging.exception(f"Fatal error: {e}")
        sys.exit(1)
